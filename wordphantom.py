import os
import requests
import argparse
import re
import docx
import math
import urllib.parse
from docx.shared import Inches
from itertools import zip_longest
from bs4 import BeautifulSoup
from urllib.parse import urlparse
from transformers import pipeline
from wordphantom.imagephantom import scrape_images

class WordPhantom:
    '''Creates summaries of scraped info from google queries'''
    def __init__(self, docx_path, batch_size=3000, num_images=1):
        self.filepath = docx_path
        self.batch_size = batch_size
        self.num_images = num_images
        self.summarizer = pipeline('summarization')
        self.links = []
        self.text_exp = ''
        self.summaries = []

    def get_links(self, query):
        '''Collects links from google'''
        self.google_url = 'https://www.google.com/search?client=ubuntu&channel=fs&q={}&ie=utf-8&oe=utf-8'.format(query)
        try:
            html = requests.get(self.google_url)
            if html.status_code == 200:
                soup = BeautifulSoup(html.text, 'lxml')
                a = soup.find_all('a')
                for i in a:
                    k = i.get('href')
                    try:
                        m = re.search("(?P<url>https?://[^\s]+)", k)
                        n = m.group(0)
                        rul = n.split('&')[0]
                        domain = urlparse(rul)
                        if(re.search('google.com', domain.netloc)):
                            continue
                        else:
                            self.links.append(rul)
                    except:
                        continue
        except Exception as ex:
            print(str(ex))
        return self.links

    def get_soup(self, url):
        '''Gets BeautifulSoup text object'''
        html = requests.get(url)
        soup = BeautifulSoup(html.text, 'lxml')
        return soup

    def get_text(self, n_links=9):
        '''Scrapes text from links'''
        BAD_URLS = {"youtube.com"}
        d = {}
        self.links = self.get_links(self.query)
        for url in self.links[:n_links]:
            url = url.split("%")[0]
            article = url.split('/')[-1]
            print(f"\n\n{url}, {d.keys()}\n{url}\n\n")

            if url[-3:] == 'pdf':
                print(f"Gross. {url} is a pdf file.")
                continue
            elif len({burl for burl in BAD_URLS if burl in url}):
                print(f"{url} in set {BAD_URLS}")
                continue
            elif len({s for s in d.keys() if article != '' \
                    and article in s.split('/')[:-1] \
                    and len(s.split('/')[:-1]) - len(article) < 6}):

                print(f"{url} similar to other in {d.keys()}")
                continue

            else:
                print(f"Getting soup for {url}")
                soup = self.get_soup(url)
                d[url] = soup.get_text()

        return d

    def get_summaries(self):
        '''Summarizes text scraped from links'''
        N = len(self.text_exp)
        print("\n\nInside get_summaries\n", N, self.text_exp)
        # maker sure n_batches is always at least 1
        n_batches = math.ceil((N+1) / self.batch_size)
        batch = N // n_batches

        for i in range(0, N, batch):
            print(i, batch+i)
            section = self.text_exp[i:(i+batch)]
            try:
                if len(section) < 50:
                    print("section too short")
                    continue

                summary = self.summarizer(section, min_length=90, max_length=200)
                self.summaries.append(summary[0]['summary_text'])
                print(summary)
            except Exception as e:
                print(f"\nFAILURE: {e}")
                continue
        return self.summaries

    def clean_summaries(self):
        '''Cleans summarized text'''
        print("Inside clean_summaries")
        self.final_text = ". ".join(sentence[0].upper() + sentence[1:] for sentence in "\n".join(self.summaries).split(" . "))
        return self.final_text

    def create_text_section(self, query):
        '''Writes MS Word Document with summarized text, pictures, and links'''
        # read or create word document and make query the heading

        self.query = query
        print("Creating document.")
        try:
            self.doc = docx.Document(self.filepath)
        except:
            self.doc = docx.Document()
        self.doc.add_heading(query, 1)
        all_summaries = []

        # scrape text
        text = self.get_text()
        print(text, type(text))
        for url, t in text.items():

            for img_path, img in scrape_images(url, n=self.num_images):
                print(f"Attempting to add img from {url}")

                try:

                    self.doc.add_picture(img_path, width=Inches(5.0))
                except Exception as e:
                    print(f"Nope: {e}")
                    continue

            try:
                self.text_exp = " ".join(line for line in t.split('\n') if len(line) > 20 or " [ " in line)
                self.summaries = self.get_summaries()
                self.final_text = self.clean_summaries()
                self.doc.add_paragraph(self.final_text)
                self.doc.save(self.filepath)

            except Exception as e:

                print(f"boo {url} ...")
                print(f"\n\nEXCEPTION: {e}\n\n")

                continue

        #self.summaries = self.get_summaries()
        #self.final_text = self.clean_summaries()
        #self.doc.add_paragraph(self.final_text)
        self.doc.add_paragraph('\n\n'.join(self.links))
        self.doc.save(self.filepath)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Scrape the googs!")
    parser.add_argument('filepath', type=str, help="Word Document full filepath")

    parser.add_argument('queries', type=str, nargs='*', help="Your google queries")
    args = parser.parse_args()
    wp = WordPhantom(args.filepath, num_images=5)
    for query in args.queries:
        wp.create_text_section(query)
    print(wp.links)
    print(wp.filepath)
    print(wp.summaries)
